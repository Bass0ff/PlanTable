import requests
import sys

from PySide2.QtCore import Qt, QSize, QDate
from PySide2.QtWidgets import *
from PySide2.QtGui import QPalette, QColor

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.shared import Pt, Mm, Cm
from datetime import datetime

server_host = "http://bass0ff.pythonanywhere.com/"
# server_host = "http://127.0.0.1:8000/"

tables = {
    "t-01": {
        'name': "Проведение открытых уроков, классных часов, предметных недель, других мероприятий",
        'fields': [
            ["Дата", "Date", (), 'date'],
            ["Предмет", "EList", ("Русский Язык", "Математика", "Алгебра", "Геометрия", "Литература", "Физика"), 'name'],
            ["Класс", "SText", (), 'studClass'],
            ["Тема", "Text", (), "theme"],
            ["Цель", "Text", (), "target"],
            ["Отметка", "Check", (), "result"]
        ],
        'pattern': {
            "id": -1,
            "type": "open_class",
            "teacher": "",
            "date": "",
            "name": "",
            "table": 1,
            "studClass": "",
            "theme": "",
            "target": "",
            "result": ""
        }
    },
    "t-02": {
        'name': "Участие в подготовке и проведении лицейских мероприятий",
        'fields': [
            ["Дата", "Date", (), "date"],
            ["Название", "EList", ("Заседание Кафедры", "Педсовет", "Педагогическое чтение", "Конференция", "Олимпиада", "Конкурс", "Выставка", "Предметная неделя"), "name"],
            ["Форма участия", "EList", ("Очная", "Заочная", "Дистанционная"), "form"],
            ["Документ", "List", ("Протокол", "Выписка", "План", "Отзыв", "Приказ"), "document"]
        ],
        'pattern': {
            "id": -1,
            "type": "organization",
            "teacher": "",
            "date": "",
            "name": "",
            "table": 2,
            "form": "",
            "document": "",
            "place": "Лицей N1"
        }
    },
    "t-03": {
        'name': "Запланированные мероприятия",
        'fields': [
            ["Мероприятие", "Text", (), "name"],
            ["Результат", "Text", (), "document"],
            ["Место проведения", "Text", (), "place"],
            ["Дата", "Date", (), "date"]
        ],
        'pattern': {
            "id": -1,
            "type": "organization",
            "teacher": "",
            "date": "",
            "name": "",
            "table": 3,
            "form": "-",
            "document": "",
            "place": ""
        }
    },
    "t-04": {
        'name': "Работа в рамках творческих групп, инновационной/стажировочной деятельности площадок",
        'fields': [
            ["Название", "Text", (), "name"],
            ["Личное участие", "Text", (), "action"],
            ["Дата", "Date", (), "date"],
            ["Результат", "Text", (), "result"]
        ],  
        'pattern': {
            "id": -1,
            "type": "expertise",
            "teacher": "",
            "date": "",
            "name": "",
            "table": 4,
            "result": "",
            "action": "",
            "level": "-"
        }
    },
    "t-05": {
        'name': "Экспертная Деятельность",
        'fields': [
            ["Дата", "Date", (), "date"],
            ["Название", "Text", (), "name"],
            ["Вид деятельности", "List", ("судья", "эксперт", "жюри"), "action"],
            ["Уровень", "List", ("отборочный", "заключительный", "дистанционный", "школьный", "муниципальный", "районный", "региональный", "всероссийский", "международный", "межмуниципальный", "межрегиональный"), "level"]
        ],
        'pattern': {
            "id": -1,
            "type": "expertise",
            "teacher": "",
            "date": "",
            "name": "",
            "table": 5,
            "result": "Проведено",
            "action": "",
            "level": ""
        }
    },
    "t-06": {
        'name': "Обучение на курсах повышения квалификации, посещение опорных школ и др.",
        'fields': [
            ["Дата", "Date", (), "date"],
            ["Тема", "Text", (), "theme"],
            ["Учреждение", "Text", (), "place"],
            ["Часы", "Number", (0, 500), "length"],
            ["Документ", "Text", (), "document"],
            ["Формат", "EList", ("очный", "дистанционный", "очный, с применением дистанционных технологий"), "form"]
        ],  
        'pattern': {
            "id": -1,
            "type": "course",
            "teacher": "",
            "date": "",
            "name": "Курс",
            "table": 6,
            "theme": "",
            "form": "",
            "document": "",
            "place": "",
            "organizer": "-",
            "length": 0
        }
    },
    "t-07": {
        'name': "Участие в сертифицированные вебинарах, семинарах и др.",
        'fields': [
            ["Дата", "Date", (), "date"],
            ["Тема", "Text", (), "theme"],
            ["Организатор", "Text", (), "organizer"],
            ["Формат", "Text", (), "form"],
            ["Часы", "Number", (0, 500), "length"],
            ["Документ", "Text", (), "document"]
        ],
        'pattern': {
            "id": -1,
            "type": "course",
            "teacher": "",
            "date": "",
            "name": "Курс",
            "table": 7,
            "theme": "",
            "form": "",
            "document": "",
            "place": "-",
            "organizer": "",
            "length": ""
        }
    },
    "t-08": {
        'name': "Участие в конкурсах профессионального мастерства",
        'fields': [
            ["Дата", "Date", (), "date"],
            ["Название", "Text", (), "name"],
            ["Уровень", "List", ("школьный", "муниципальный", "региональный", "всероссийский"), "level"],
            ["Формат", "List", ("очная", "заочная", "дистанционная"), "form"],
            ["Этап", "List", ("отборочный", "заключительный", "дистанционный", "школьный", "муниципальный", "районный", "региональный", "всероссийский", "международный", "межмуниципальный", "межрегиональный"), "place"],
            ["Результат", "Text", (), "result"],
            ["Документ", "Text", (), "document"],
            ["Ссылка", "Text", (), "link"]
        ],
        'pattern': {
            "id": -1,
            "type": "experience",
            "teacher": "",
            "date": "",
            "name": "",
            "table": 8,
            "theme": "-",
            "result": "",
            "form": "",
            "document": "",
            "place": "",
            "action": "-",
            "level": "",
            "link": ""
        }
    },
    "t-09": {
        'name': "Обобщение и представление опыта работы",
        'fields': [
            ["Дата", "Date", (), "date"],
            ["Название", "Text", (), "name"],
            ["Форма участия", "List", ("очная", "заочная", "дистанционная"), "form"],
            ["Уровень", "List", ("отборочный", "заключительный", "дистанционный", "школьный", "муниципальный", "районный", "региональный", "всероссийский", "международный", "межмуниципальный", "межрегиональный"), "level"],
            ["Тема", "Text", (), "theme"],
            ["Вид деятельности", "EList", ("выступление", "публикация", "мастер-класс"), "action"],
            ["Публикация", "EList", ("статья", "метод", "разработка"), "document"],
            ["Орган", "Text", (), "place"],
            ["Ссылка", "Text", (), "link"]
        ],
        'pattern': {
            "id": -1,
            "type": "experience",
            "teacher": "",
            "date": "",
            "name": "",
            "table": 9,
            "theme": "",
            "result": "-",
            "form": "",
            "document": "",
            "place": "",
            "action": "",
            "level": "",
            "link": ""
        }
    },
    "t-10": {
        'name': "Участие в диагностике профессиональных дефицитов/предметных компетенций",
        'fields': [
            ["Дата", "Date", (), "date"],
            ["Название", "Text", (), "name"],
            ["Результат", "Text", (), "result"]
        ],
        'pattern': {
            "id": -1,
            "type": "expertise",
            "teacher": "",
            "date": "",
            "name": "",
            "table": 10,
            "result": "",
            "action": "-",
            "level": "-"
        }
    },
    "t-11": {
        'name': "Участие во внешкольных мероприятий",
        'fields': [
            ["Дата", "Date", (), "date"],
            ["Тип", "EList", ("Школа Современного Педагога", "конференция", "семинар", "консультация", "стажировочная площадка", "урок коллег из другой школы"), "name"],
            ["Уровень", "List", ("отборочный", "заключительный", "дистанционный", "школьный", "муниципальный", "районный", "региональный", "всероссийский", "международный", "межмуниципальный", "межрегиональный"), "level"],
            ["Статус", "List", ("организатор", "участник"), "form"],
            ["Место проведения", "Text", (), "place"],
            ["Тема", "Text", (), "theme"],
            ["Организатор", "Text", (), "organizer"]
        ],
        'pattern': {
            "id": -1,
            "type": "expertise",
            "teacher": "",
            "date": "",
            "name": "",
            "table": 11,
            "theme": "",
            "form": "",
            "document": "",
            "place": "",
            "organizer": "",
            "length": 0
        }
    },
    "t-12": {
        'name': "Посещение уроков, кл.часов, мероприятий у коллег в школе",
        'fields': [
            ["Дата", "Date", (), "date"],
            ["Предмет", "EList", ("Русский Язык", "Математика", "Алгебра", "Геометрия", "Литература", "Физика"), "name"],
            ["Класс", "SText", (), "studClass"],
            ["Тема", "Text", (), "theme"],
            ["Цель", "Text", (), "target"]
        ],
        'pattern': {
            "id": -1,
            "type": "open_class",
            "teacher": "",
            "date": "",
            "name": "",
            "table": 12,
            "studClass": "",
            "theme": "",
            "target": "",
            "result": "Посещено"
        }
    },
    "t-13": {
        'name': "Участие обучающихся в конкурсных мероприятиях, входящих в перечень, \n утвержденный приказом Министертсва науки и высшего образования РФ",
        'fields': [
            ["Дата", "Date", (), "date"],
            ["Название", "Text", (), "name"],
            ["Обучающийся", "Text", (), "student"],
            ["Класс", "SText", (), "studClass"],
            ["Этап", "EList", ("отборочный", "заключительный", "дистанционный", "школьный", "муниципальный", "районный", "региональный", "всероссийский", "международный", "межмуниципальный", "межрегиональный"), "level"],
            ["Результат", "Text", (), "result"],
            ["Документ", "Text", (), "document"]
        ],
        'pattern': {
            "id": -1,
            "type": "student_work",
            "teacher": "",
            "date": "",
            "name": "",
            "table": 13,
            "result": "",
            "theme": "-",
            "student": "",
            "studClass": "",
            "level": "",
            "document": ""
        }
    },
    "t-14": {
        'name': "Участие обучающихся в других конкурсных мероприятиях, научно-практических конференциях, ШРД, ФНР и др.",
        'fields': [
            ["Дата", "Date", (), "date"],
            ["Название", "Text", (), "name"],
            ["Обучающийся", "Text", (), "student"],
            ["Класс", "SText", (), "studClass"],
            ["Уровень", "EList", ("отборочный", "заключительный", "дистанционный", "школьный", "муниципальный", "районный", "региональный", "всероссийский", "международный", "межмуниципальный", "межрегиональный"), "level"],
            ["Результат", "Text", (), "result"],
            ["Документ", "Text", (), "document"]
        ],
        'pattern': {
            "id": -1,
            "type": "student_work",
            "teacher": "",
            "date": "",
            "name": "",
            "table": 14,
            "result": "",
            "theme": "-",
            "student": "",
            "studClass": "",
            "level": "",
            "document": ""
        }
    },
    "t-15": {
        'name': "Участие обучающихся в соревнованиях профессиональных компетенций",
        'fields': [
            ["Дата", "Date", (), "date"],
            ["Название", "Text", (), "name"],
            ["Компетенция", "Text", (), "theme"],
            ["Обучающийся", "Text", (), "student"],
            ["Класс", "SText", (), "studClass"],
            ["Уровень", "EList", ("отборочный", "заключительный", "дистанционный", "школьный", "муниципальный", "районный", "региональный", "всероссийский", "международный", "межмуниципальный", "межрегиональный"), "level"],
            ["Результат", "Text", (), "result"],
            ["Документ", "Text", (), "document"]
        ],
        'pattern': {
            "id": -1,
            "type": "student_work",
            "teacher": "",
            "date": "",
            "name": "",
            "table": 15,
            "result": "",
            "theme": "",
            "student": "",
            "studClass": "",
            "level": "",
            "document": ""
        }
    },
    "t-16": {
        'name': "Дополнительные общеразвивающие программы (ДОП) по подготовке обучющихся 9-11 классов к ВсОШ",
        'fields': [
            ["Название", "Text", (), "name"],
            ["Дата", "Date", (), "date"],
            ["Обучающийся", "Text", (), "student"]
        ],
        'pattern': {
            "id": -1,
            "type": "student_work",
            "teacher": "",
            "date": "",
            "name": "Курс",
            "table": 16,
            "result": "-",
            "theme": "-",
            "student": "",
            "studClass": "-",
            "level": "-",
            "document": "-"
        }
    },
    "t-17": {
        'name': "Участие в профильных сменах",
        'fields': [
            ["Название", "Text", (), "name"],
            ["Дата", "Date", (), "date"],
            ["Обучающийся", "Text", (), "student"]
        ],
        'pattern': {
            "id": -1,
            "type": "student_work",
            "teacher": "",
            "date": "",
            "name": "Курс",
            "table": 17,
            "result": "-",
            "theme": "-",
            "student": "",
            "studClass": "-",
            "level": "-",
            "document": "-"
        }
    }
}

docTables = [
    [   "Учебно-методическая и организационно-методическая работа",
        ["Проведение открытых уроков, классных часов, предметных недель, других мероприятий", ["Дата", "Предмет", "Класс", "тема", "Цель, для какой цели проводится", "Отметка о выполнении"], ["date", "name", "studClass", "theme", "target", "result"]],
        ["Участие в подготовке и проведении лицейских мероприятий", ["Дата", "Название мероприятия", "Форма участия", "Вид сданной документации"], ["date", "name", "form", "document"]],
    ],
    [   "Научно-методическая и исследовательская (экспериментальная) работа",
        ["Тема самообразования", ["Запланированные мероприятия", "Конкретный результат", "Место проведения", "Дата"], ["name", "document", "place", "date"]],
        ["Работа в рамках творческих групп, инновационной/стажировочной деятельности площадок", ["Название творческой группы, инновационной/стажировочной площадки", "Личное участие в работе группы, площадки", "Дата", "Результат"], ["name", "action", "date", "result"]],
        ["Экспертная Деятельность", ["Дата", "Название мероприятия", "Вид экспертной детельности", "Уровень"], ["date", "name", "action", "level"]],
    ],
    [   "Повышение квалификации",
        ["Обучение на курсах повышения квалификации, посещение опорных школ и др.", ["Дата обучения", "Тема курсовой подготовки", "Базовое учреждение обучения (по удостоверению)", "Количество часов", "Документ об окончании обучения", "Формат обучения"], ["date", "theme", "place", "length", "document", "form"]],
        ["Участие в сертифицированные вебинарах, семинарах и др.", ["Дата", "Тема мероприятия", "Организатор мероприятия", "Формат обучения", "Количество часов", "Документ"], ["date", "theme", "organizer", "form", "length", "document"]],
        ["Участие в конкурсах профессионального мастерства", ["Дата", "Название", "Уровень", "Формат", "Этап", "Результат участия", "Документ", "Активная ссылка на размещение материалов в сети интернет"], ["date", "name", "level", "form", "place", "result", "document", "link"]],
        ["Обобщение и представление опыта работы", ["Дата", "Название мероприятия", "Форма участия", "Уровень", "Тема представления опыта", "Выступление, публикация, мастер-класс", "Вид публикации", "Название органа, издания, исходные данные", "Активная ссылка на размещение материалов в сети интернет"], ["date", "name", "form", "level", "theme", "action", "document", "result", "link"]],
        ["Участие в диагностике профессиональных дефицитов/предметных компетенций", ["Дата", "Название диагностики", "Результат"], ["date", "name", "result"]],
        ["Участие во внешкольных мероприятий", ["Дата", "Тип мероприятия", "Уровень", "Статус", "Место проведения", "Тема мероприятия", "Кто проводил"], ["date", "name", "document", "action", "place", "theme", "organizer"]],
        ["Посещение уроков, кл.часов, мероприятий у коллег в школе", ["Дата", "Предмет", "Класс", "Тема", "Цель проведения, для какой категории проводится"], ["date", "name", "studClass", "theme", "target"]],
    ],
    [   "Работа с обучающимися, в том числе и внеучебная",
        ["Участие обучающихся в конкурсных мероприятиях, входящих в перечень, утвержденный приказом Министерcтва науки и высшего образования РФ", ["Дата", "Наименование мероприятия", "ФИ обучающегося, класс", "Этап олимпиады", "Результат уастия, подтверждающий документ"], ["date", "name", "student", "level", "result"]],
        ["Участие обучающихся в других конкурсных мероприятиях, научно-практических конференциях, ШРД, ФНР и др.", ["Дата", "Наименование мероприятия", "ФИ обучающегося или группа учеников, класс", "Уровень", "Результат участия, подтверждающий документ"], ["date", "name", "student", "level", "result"]],
        ["Участие обучающихся в соревнованиях профессиональных компетенций", ["Дата", "Наименование соревнований", "Название компетенции", "ФИ обучающегося или группа учеников, класс", "Уровень", "Результат участия, подтверждающий документ"], ["date", "name", "theme", "student", "level", "result"]],
        ["Участие обучающихся в программах образовательного фонда «Талант и успех» (образовательные центры «Сириус» и «Персей»)", ["Название программы", "Сроки прохождения", "ФИ обучающегося", "Название смены", "Сроки", "ФИ обучающегося (участника смены)"], ["name1", "date1", "student1", "name2", "date2", "student2"]]
    ]
]

def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def make_row_bold(row: tuple):
    for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

class QHLine(QFrame):
    def __init__(self):
        super(QHLine, self).__init__()
        self.setFrameShape(QFrame.HLine)
        self.setFrameShadow(QFrame.Sunken)

class RowForm(QDialog):
    def __init__(self, root, data:dict = {}):
        self.root = root
        super().__init__()
        self.setWindowTitle(f'Форма "{tables[root.name]["name"]}"')
        self.setWindowModality(Qt.WindowModality.ApplicationModal)
        self.setMinimumWidth(800)
        self.setMinimumHeight(600)
        pattern = tables[root.name]["fields"]
        self.layout = QVBoxLayout()

        for i in range(len(pattern)):
            row = QHBoxLayout()
            row.setAlignment(Qt.AlignTop | Qt.AlignLeft)
            row.addWidget(QLabel(pattern[i][0]), stretch=1)
            field = 0
            db_f = pattern[i][3]
            match pattern[i][1]:
                case "Text":
                    field = QTextEdit()
                    field.setMaximumWidth(640)
                    field.setMinimumHeight(30)
                    field.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
                    try:
                        field.setText(data[db_f])
                    except KeyError:
                        pass
                case "SText":
                    field = QLineEdit()
                    field.setMaximumWidth(50)
                    field.setMinimumHeight(30)
                    field.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Preferred)
                    field.setMaxLength(3)
                    try:
                        field.setText(data[db_f])
                    except KeyError:
                        pass
                case "Number":
                    field = QSpinBox()
                    field.setMinimum(pattern[i][2][0])
                    field.setMaximum(pattern[i][2][1])
                    field.setMaximumWidth(50)
                    field.setMinimumHeight(30)
                    field.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Preferred)
                    try:
                        field.setValue(data[db_f])
                    except KeyError:
                        pass
                case "Date":
                    field = QDateEdit()
                    field.setDate(QDate.currentDate())
                    field.setMinimumHeight(30)
                    field.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Preferred)
                    try:
                        field_data = data[db_f]
                        if '-' in field_data:
                            field_data = data[db_f].split("-")
                        elif '.' in field_data:
                            field_data = data[db_f].split(".")
                        field.setDate(QDate(int(field_data[2]),int(field_data[1]),int(field_data[0])))
                    except KeyError:
                        pass
                case "List":
                    field = QComboBox()
                    field.addItems(pattern[i][2])
                    field.setMinimumHeight(30)
                    field.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Preferred)
                    try:
                        index = field.findText(data[db_f])
                        field.setCurrentIndex(index)
                    except KeyError:
                        pass
                case "EList":
                    field = QComboBox()
                    field.setEditable(True)
                    field.addItems(pattern[i][2])
                    try:
                        index = field.findText(data[db_f])
                        if index == -1:
                            field.addItem(data[db_f])
                        index = field.findText(data[db_f])
                        field.setCurrentIndex(index)  
                    except KeyError:
                        pass
                case "Check":
                    field = QCheckBox()
                    field.setFixedHeight(30)
                    field.setMinimumHeight(30)
                    field.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Preferred)
                    try:
                        if data[db_f] == "True":
                            field.setChecked(True)
                    except KeyError:
                        pass
            row.addWidget(field, stretch=4)
            self.layout.addLayout(row)
        if self.root.win.access != "Учитель" and not ('teacher' in data.keys()):
            data['teacher'] = self.root.win.teach
        if self.root.win.access != "Учитель" and not ('teacherName' in data.keys()):
            data['teacherName'] = self.root.win.name
        self.data = data

        if self.root.win.access != "Учитель" and len(data) > 0:
            row = QHBoxLayout()
            row.setAlignment(Qt.AlignTop | Qt.AlignLeft)
            row.addWidget(QLabel("Преподаватель:"), stretch=1)
            row.addWidget(QLabel(data['teacherName']), stretch=4)
            self.layout.addLayout(row)
        
        self.layout.addWidget(QHLine())
        menu = QHBoxLayout()
        save = QPushButton("ОК")
        save.clicked.connect(self.save)
        abort = QPushButton("Отмена")
        abort.clicked.connect(self.abort)
        delete = QPushButton("Удалить")
        delete.clicked.connect(self.delete)
        menu.addWidget(save)
        menu.addWidget(abort)
        menu.addWidget(delete)
        self.layout.addLayout(menu)
        self.setLayout(self.layout)

    def data_format(self, line):
        a = type(line)
        if a == QDateEdit:
            return str(line.date().day()) + "." + str(line.date().month()) + "." + str(line.date().year())
        elif a == QLineEdit:
            return line.text()
        elif a == QTextEdit:
            return line.toPlainText()
        elif a == QComboBox:
            return line.currentText()
        elif a == QCheckBox:
            return str(line.isChecked())
        elif a == QSpinBox:
            return line.value()

    def save(self):
        if self.root.win.access != "Учитель" and int(self.data['teacher']) != int(self.root.win.teach):
            print(f"НЕЛЬЗЯ МЕНЯТЬ ЧУЖИЕ ДАННЫЕ! ID ПОЛЬЗОВАТЕЛЯ: {self.root.win.teach}, ID АВТОРА: {self.data['teacher']}")
            self.root.flag = "ABORT"
        else:
            rowData = tables[self.root.name]['pattern'].copy()
            rowData['teacher'] = self.root.win.teach
            if self.root.win.access != "Учитель":
                add = 1
            else:
                add = 0
            for i in range(self.layout.count()-2-add):
                val = self.data_format(self.layout.itemAt(i).layout().itemAt(1).widget())
                name = self.layout.itemAt(i).layout().itemAt(0).widget().text()
                field = tables[self.root.name]['fields'][i][3]
                if self.root.flag == "NEW":
                    rowData[field] = val
                else:
                    self.data[field] = val
                if (name == "Тип" or name == "Название" or name == "Мероприятие" or name == "Дата" or name == "Предмет" or name == "Класс"):
                    valab = QLabel(val)
                    valab.setStyleSheet('border: 1px solid black;')
                    if (name == "Дата"):
                        if self.root.flag == "EDIT":
                            self.root.curRow.itemAt(0).widget().setText(val)
                        else:
                            self.root.curRow.insertWidget(0, valab, stretch=2)
                    elif name == "Класс":
                        if self.root.flag == "EDIT":
                            self.root.curRow.itemAt(2).widget().setText(val)
                        else:
                            self.root.curRow.insertWidget(2, valab, stretch=1)
                    else:
                        if self.root.flag == "EDIT":
                            self.root.curRow.itemAt(1).widget().setText(val)
                        else:
                            self.root.curRow.insertWidget(1, valab, stretch=4)
            if self.root.flag == "NEW":
                self.root.data.append(rowData)
            
        
            self.root.flag = "OK"
        self.close()

    def abort(self):
        self.root.flag = "ABORT"
        self.close()
    
    def delete(self):
        self.root.flag = "DEL"
        self.close()

class AuthDialog(QDialog):
    def __init__(self, root):
        self.root = root
        super().__init__()
        self.user = ""
        self.setWindowTitle(f'Авторизация')
        self.setWindowModality(Qt.WindowModality.ApplicationModal)
        self.layout = QVBoxLayout()
        row = QHBoxLayout()
        lbl = QLabel("Пользователь")
        row.addWidget(lbl, stretch=2)
        self.user = QLineEdit()
        row.addWidget(self.user, stretch=3)
        self.layout.addLayout(row)
        row = QHBoxLayout()
        lbl = QLabel("Пароль")
        row.addWidget(lbl, stretch=2)
        self.password = QLineEdit()
        self.password.setEchoMode(QLineEdit.Password)
        row.addWidget(self.password, stretch=3)
        self.layout.addLayout(row)

        self.layout.addWidget(QHLine())
        menu = QHBoxLayout()
        OK = QPushButton("ОК")
        OK.clicked.connect(self.check)
        abort = QPushButton("Закрыть")
        abort.clicked.connect(self.cancel)
        reg = QPushButton("Новый профиль")
        reg.clicked.connect(self.reg)
        ps = QPushButton("<O>")
        ps.clicked.connect(self.togglePass)
        
        menu.addWidget(OK)
        menu.addWidget(abort)
        menu.addWidget(reg)
        menu.addWidget(ps)

        self.layout.addLayout(menu)

        self.setLayout(self.layout)

    def togglePass(self):
        if self.password.echoMode() == QLineEdit.Password:
            self.password.setEchoMode(QLineEdit.Normal)
        else:
            self.password.setEchoMode(QLineEdit.Password)


    def check(self):
        u = self.user.text()
        p = self.password.text()
        response = requests.get(server_host + "auth", params={"user": u, "pass": p})
        if response.text == "NO":
            err = QMessageBox()
            err.setText("Пароль или имя пользователя введено неверно.")
            err.setInformativeText("Проверьте правильность пароля и имени и попробуйте ещё раз.")
            err.exec_()
        else:
            self.root.flag = f"{u}, {response.text}"
            self.close()
       
    def cancel(self):
        self.root.flag = "NVM"
        self.close()

    def reg(self):
        self.flag = "NVM"
        r = RegDialog(self)
        r.exec_()
        if self.flag != "NVM":
            self.root.flag = self.flag
            self.close()

class RegDialog(QDialog):
    def __init__(self, root):
        self.root = root
        super().__init__()
        self.setWindowTitle(f'Регистрация')
        self.setWindowModality(Qt.WindowModality.ApplicationModal)
        self.layout = QVBoxLayout()

        row = QHBoxLayout()
        lbl = QLabel("ФИО")
        row.addWidget(lbl, stretch=1)
        self.user = QLineEdit()
        row.addWidget(self.user, stretch=3)
        self.layout.addLayout(row)
        row = QHBoxLayout()
        lbl = QLabel("Предмет")
        row.addWidget(lbl, stretch=1)
        self.subj = QLineEdit()
        row.addWidget(self.subj, stretch=3)
        self.layout.addLayout(row)
        row = QHBoxLayout()
        lbl = QLabel("Категория")
        row.addWidget(lbl, stretch=1)
        self.category = QLineEdit()
        row.addWidget(self.category, stretch=3)
        self.layout.addLayout(row)
        row = QHBoxLayout()
        lbl = QLabel("Уровень доступа")
        row.addWidget(lbl, stretch=1)
        self.access = QComboBox()
        self.access.addItems(['Учитель', 'Зав. кафедрой', 'Методист'])
        row.addWidget(self.access, stretch=3)
        self.layout.addLayout(row)

        row = QHBoxLayout()
        lbl = QLabel("Кафедра")
        row.addWidget(lbl, stretch=1)
        self.department = QComboBox()
        self.department.addItems(['Кафедра точных наук', 'Кафедра естественных наук', 'Кафедра общественно-гуманитарных наук', 'Метод. объединение учителей физкультуры'])
        row.addWidget(self.department, stretch=3)
        self.layout.addLayout(row)

        row = QHBoxLayout()
        lbl = QLabel("Пароль")
        row.addWidget(lbl, stretch=1)
        self.password = QLineEdit()
        self.password.setEchoMode(QLineEdit.Password)
        row.addWidget(self.password, stretch=3)
        self.layout.addLayout(row)
        row = QHBoxLayout()
        lbl = QLabel("Повтор пароля")
        row.addWidget(lbl, stretch=1)
        self.passrep = QLineEdit()
        self.passrep.setEchoMode(QLineEdit.Password)
        row.addWidget(self.passrep, stretch=3)
        self.layout.addLayout(row)

        self.layout.addWidget(QHLine())
        menu = QHBoxLayout()
        OK = QPushButton("ОК")
        OK.clicked.connect(self.save)
        abort = QPushButton("Отмена")
        abort.clicked.connect(self.abort)
        ps = QPushButton("<O>")
        ps.clicked.connect(self.togglePass)
        
        menu.addWidget(OK)
        menu.addWidget(abort)
        menu.addWidget(ps)

        self.layout.addLayout(menu)

        self.setLayout(self.layout)

    def togglePass(self):
        if self.password.echoMode() == QLineEdit.Password:
            self.password.setEchoMode(QLineEdit.Normal)
            self.passrep.setEchoMode(QLineEdit.Normal)
        else:
            self.password.setEchoMode(QLineEdit.Password)
            self.passrep.setEchoMode(QLineEdit.Password)

    def save(self):
        if self.password.text() != self.passrep.text():
            dlg = QMessageBox()
            dlg.setWindowTitle("Ошибка")
            dlg.setText("Введённый пароль не совпадает с полем подтверждения!")
            dlg.exec_()
            return 0

        name = self.user.text()
        subj = self.subj.text()
        ctgr = self.category.text()
        acss = self.access.currentText()
        pwrd = self.password.text()
        dprt = self.department.currentText()
        response = requests.get(server_host + "reg", params={"name": name, "pass": pwrd, "subj": subj, "category": ctgr, "access": acss, "department": dprt})
        if response.text == "AE":
            err = QMessageBox()
            err.setText("Регистрация невозможна.")
            err.setInformativeText("Преподаватель данного предмета с таким именем уже зарегистрирован.")
            err.exec_()
        else:
            self.root.flag = f"{name}, {response.text}, {subj}, {acss}, {ctgr}, {dprt}"
            self.close()

    def abort(self):
        self.root.flag = "NVM"
        self.close()

class DocDialog(QDialog):
    def __init__(self, root):
        self.root = root
        super().__init__()
        self.setWindowTitle(f'Регистрация')
        self.setWindowModality(Qt.WindowModality.ApplicationModal)
        self.layout = QVBoxLayout()

        genBtn = QPushButton("Отчёт по школе")
        genBtn.clicked.connect(lambda: self.done(3))
        self.layout.addWidget(genBtn)
        if self.root.access != "Методист" and self.root.access != "dev":
            genBtn.setEnabled(False)
        depBtn = QPushButton("Отчёт по кафедре")
        depBtn.clicked.connect(lambda: self.done(2))
        self.layout.addWidget(depBtn)
        if self.root.access == "Учитель":
            depBtn.setEnabled(False)
        indBtn = QPushButton("Индивидуальный отчёт")
        indBtn.clicked.connect(lambda: self.done(1))
        self.layout.addWidget(indBtn)
        canBtn = QPushButton("Отмена")
        canBtn.clicked.connect(lambda: self.done(0))
        self.layout.addWidget(canBtn)

        self.setLayout(self.layout)

class TableBtn(QPushButton):
    def __init__(self, *a, **kw):
        super(TableBtn, self).__init__(*a, **kw)
    
    def index(self, id: int):
        self.id = id

    def enterEvent(self, event):
        self.old = self.styleSheet()
        self.setStyleSheet('border: 1px solid blue;\
                                background-color: #9999ff')
        return super(TableBtn, self).enterEvent(event)

    def leaveEvent(self, event):
        self.setStyleSheet(self.old)
        return super(TableBtn, self).enterEvent(event)

class Table(QWidget):
    def __init__(self, win, table_name: str = ""):
        self.name = table_name
        self.win = win
        self.data = []
        self.delData = []
        self.draw()
        
    def draw(self):
        self.widget = QWidget()
        layout = QHBoxLayout()
        layout.setAlignment(Qt.AlignLeft)
        
        Box_Left = QVBoxLayout()
        Box_Left.setAlignment(Qt.AlignTop)
        Box_Right = QVBoxLayout()
        Box_Right.setAlignment(Qt.AlignTop)

        left_widget = Color('#80B9C7')
        left_widget.setMaximumWidth(self.win.width()/7*3)
        left_widget.setMinimumHeight(self.win.height())
        left_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        left_widget.setLayout(Box_Left)

        right_widget = Color('white')
        right_widget.setMinimumWidth(self.win.width()/5*3)
        right_widget.setMinimumHeight(self.win.height())
        right_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        right_widget.setLayout(Box_Right)
        
        btn_menu = QPushButton('Учебно-методическая и \nорганизационно-методическая работа ')
        btn_menu.clicked.connect(lambda: self.win.goTo(self.win.page_plan_p1))
        Box_Left.addWidget(btn_menu)

        btn_menu = QPushButton('Научно-методическая \nи исследовательская работа')
        btn_menu.clicked.connect(lambda: self.win.goTo(self.win.page_plan_p2))
        Box_Left.addWidget(btn_menu)

        btn_menu = QPushButton('Повышение квалификации')
        btn_menu.clicked.connect(lambda: self.win.goTo(self.win.page_plan_p3))
        Box_Left.addWidget(btn_menu)

        btn_menu = QPushButton('Работа с обучающимися, \nв том числе и внеучебная')
        btn_menu.clicked.connect(lambda: self.win.goTo(self.win.page_plan_p4))
        Box_Left.addWidget(btn_menu)

        btn_menu = QPushButton('К созданию документа')
        btn_menu.clicked.connect(self.win.docx_gen)
        Box_Left.addWidget(btn_menu)

        if self.name == "t-03":
            texts = QFormLayout()
            self.selfed_id = -1
            self.theme = QLineEdit()
            self.stage = QLineEdit()
            self.method = QLineEdit()
            texts.addRow("Тема Самообразования", self.theme)
            texts.addRow("Сроки работы над темой (год, этап)", self.stage)
            texts.addRow("Образовательная технология/Метод обучения", self.method)
            texts.setAlignment(Qt.AlignBottom)
            Box_Right.addLayout(texts)

        lbl = QLabel(tables[self.name]["name"])
        font = lbl.font()
        font.setBold(True)
        font.setPointSize(14)
        lbl.setFont(font)
        lbl.setAlignment(Qt.AlignHCenter | Qt.AlignBottom)
        lbl.setWordWrap(True)
        lbl.setFixedHeight(110)
        Box_Right.addWidget(lbl)

        BoxPrev = QWidget()
        prev = QHBoxLayout(BoxPrev)
        prev.setAlignment(Qt.AlignRight)
        self.btn_prev = QPushButton("<- Предыдущая Таблица")
        self.btn_prev.clicked.connect(lambda: self.win.prev_plan_table(self))
        prev.addWidget(self.btn_prev)
        Box_Right.addWidget(BoxPrev)

        BoxTable = Color('#bababa')
        table_box = QVBoxLayout(BoxTable)
        table_box.setAlignment(Qt.AlignTop | Qt.AlignLeft)
        
        tableWidget = QWidget()
        tableWidget.setMinimumSize(self.win.width()/5 * 2, self.win.height()/ 4)
        self.rows = QVBoxLayout(tableWidget)
        self.rows.setAlignment(Qt.AlignTop | Qt.AlignLeft)
        self.rows.setContentsMargins(0, 5, 10, 5)
        self.rows.setSizeConstraint(QLayout.SetMinAndMaxSize)

        newRowBtn = TableBtn("Добавить строку")
        newRowBtn.clicked.connect(self.new_row)
        newRowBtn.setStyleSheet('border: 1px solid black;\
                                background-color: #ffffff')
        newRowBtn.setMinimumHeight(50)
        newRowBtn.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        self.rows.addWidget(newRowBtn)

        params={"id": self.win.teach, "table": int(self.name[-2:]), "type": tables[self.name]['pattern']['type']}
        response = requests.get(server_host + "getData", params=params)

        if self.name == "t-03":
            if response.json()['selfEd'] != "None":
                self.selfed_id = response.json()['selfEd']['id']
                self.method.setText(response.json()['selfEd']['method'])
                self.theme.setText(response.json()['selfEd']['theme'])
                self.stage.setText(response.json()['selfEd']['stage'])

        lines = response.json()['data']
        if len(lines) > 0:
            print(f"ТАБЛИЦА {self.name}: ДАННЫЕ С СЕРВЕРА:")
            for line in lines:
                print(line)
                self.new_row(line)
            print()

        scroller = QScrollArea()
        scroller.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOn)
        scroller.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroller.setWidgetResizable(True)                           
        scroller.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        scroller.setAlignment(Qt.AlignTop | Qt.AlignLeft)
        scroller.setWidget(tableWidget)
        table_box.addWidget(scroller)

        save = QWidget()
        BoxSave = QHBoxLayout(save)
        SaveBtn = QPushButton("СОХРАНИТЬ")
        SaveBtn.setAutoFillBackground(True)
        SaveBtn.setStyleSheet("background-color: #11aa00;")
        palette = SaveBtn.palette()
        palette.setColor(QPalette.Window, QColor('green'))
        SaveBtn.setPalette(palette)
        font = SaveBtn.font()
        font.setPointSize(16)
        font.setBold(True)
        SaveBtn.setFont(font)
        SaveBtn.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
        SaveBtn.clicked.connect(self.db_save)
        BoxSave.addWidget(SaveBtn)
        BoxSave.setAlignment(Qt.AlignHCenter)

        Box_Right.addWidget(BoxTable, stretch=3)
        Box_Right.addWidget(save)

        BoxNext = QWidget()
        next = QHBoxLayout(BoxNext)
        next.setAlignment(Qt.AlignRight)
        self.Btn_next = QPushButton("Следующая Таблица ->")
        self.Btn_next.clicked.connect(lambda: self.win.next_plan_table(self))
        next.addWidget(self.Btn_next)
        Box_Right.addWidget(BoxNext)
        
        layout.addWidget(left_widget, stretch=2)
        layout.addWidget(right_widget, stretch=3)       
        self.widget.setLayout(layout)

    def db_save(self):
        print(f"СОХРАНЯЕМ ТАБЛИЦУ {self.name}:")
        se_flag = ""
        if self.name == "t-03":
            print(f"  -> Сохраняем тему самообразования...")
            response = requests.get(server_host + "upData", params = {"type": "selfEd", "id": self.selfed_id, "teacher": self.win.teach, "method": self.method.text(), "theme": self.theme.text(), "stage": self.stage.text()})
            if response.json()['text'] == "Updated":
                print(f"    -> Обновлена запись по самообразованию")
                se_flag = "updated"
            elif response.json()['text'] == "Created":
                print(f"    -> Создана запись по новой теме самообразования")
                self.selfed_id = int(response.json()['id'])
                se_flag = "created"
        dd_cnt = 0
        if self.delData:
            print(f"  -> Удаляем стёртые записи...")
            for i in range(len(self.delData)):
                response = requests.get(server_host + "unData", params={"id": self.delData.pop()})
                print(f"    -> Удалено событие {response.text}")
                dd_cnt += 1

        ad_cnt_n = 0
        ad_cnt_o = 0
        print(f"  -> Сохраняем локальные записи...")
        for i in range(len(self.data)):
            if int(self.data[i]['teacher']) != int(self.win.teach):
                print("    ПРОПУСКАЕМ ЧУЖУЮ ЗАПИСЬ")
                continue
            pack = self.data[i]
            pack['teacher'] = self.win.teach
            response = requests.get(server_host + "upData", params=pack)
            print(pack)
            if response.text == "Ok":
                print(f"    -> Сохранена запись о мероприятии за {pack['date']}")
                ad_cnt_o += 1
            elif response.json()['text'] == "Newline":
                print(f"    -> Сохранена запись о новом мероприятии за {pack['date']}")
                self.data[i]["id"] = response.json()['id']
                ad_cnt_n += 1
        print("Дело сделано!")

        message = f'Сохранение данных завершено!'
        if se_flag == "updated":
            message += f"\nОбновлена запись по самообразованию"
        elif se_flag == "created":
            message += f"\nСоздана новая запись по самообразованию"

        if dd_cnt > 0:
            message += f"\nУдалено {dd_cnt} записей."

        if ad_cnt_o > 0:
            message += f"\nОбновлено {ad_cnt_o} записей."

        if ad_cnt_n > 0:
            message += f"\nДобавлено {ad_cnt_n} записей."

        dlg = QMessageBox()
        dlg.setWindowTitle("Сохранение")
        dlg.setText(message)
        dlg.exec_()

    def new_row(self, data={}):
        self.flag = "NEW"
        if not data:
            data = {}
        newRow = TableBtn()
        newRow.setStyleSheet('border: 1px solid black;\
                                background-color: #ffffff')
        newRow.setMinimumHeight(50)
        newRow.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Preferred)
        newRow.clicked.connect(lambda: self.edit_row(self.rows.indexOf(newRow)))
        self.curRow = QHBoxLayout(newRow)

        if len(data) == 0:
            self.f = RowForm(self)
            self.f.exec_()

            if self.flag == "OK":
                self.rows.insertWidget(self.rows.count()-1, newRow)
        else:
            rowData = tables[self.name]['pattern'].copy()
            rowData['teacher'] = self.win.teach
            if self.win.access != "Учитель":
                rowData['teacherName'] = data['teacherName']
            if int(rowData['teacher']) != int(data['teacher']):
                newRow.setStyleSheet('border: 1px solid gray;\
                                background-color: #dddddd')
            for key in rowData.keys():
                rowData[key] = data[key]
                if (key == "date"):
                        formDate = data[key].split("-")
                        formDate = f"{formDate[2]}.{formDate[1]}.{formDate[0]}"
                        rowData[key] = formDate
        
            valab = QLabel(rowData['date'])
            valab.setStyleSheet('border: 1px solid black;')
            self.curRow.insertWidget(0, valab, stretch=2)
            valab = QLabel(rowData['name'])
            valab.setStyleSheet('border: 1px solid black;')
            self.curRow.insertWidget(1, valab, stretch=4)
            if "studClass" in rowData.keys():
                valab = QLabel(rowData['studClass'])
                valab.setStyleSheet('border: 1px solid black;')
                self.curRow.insertWidget(2, valab, stretch=1)
               
                
            self.rows.insertWidget(self.rows.count()-1, newRow)
            self.data.append(rowData)
            
    def edit_row(self, index):
        print(index, end="; ")
        print(self.data[index]["id"])
        self.flag = "EDIT"
        self.curRow = self.rows.itemAt(index).widget().children()[0]

        self.f = RowForm(self, self.data[index])
        self.f.exec()

        if self.flag == "DEL":
            if self.data[index]['id'] >= 0:
                self.delData.append(self.data[index]['id'])
            else:
                print("Строка не была сохранена в БД. Удаление локальной записи.")
            self.rows.itemAt(index).widget().deleteLater()
            self.data.pop(index)
        ...

class Color(QWidget):
    def __init__(self, color):
        super(Color, self).__init__()
        self.setAutoFillBackground(True)
        palette = self.palette()
        palette.setColor(QPalette.Window, QColor(color))
        self.setPalette(palette)

class PDataForm(QWidget):
    def __init__(self, root, data: dict = {}):
        super(PDataForm, self).__init__()
        self.data = data
        self.root = root
        self.pattern = [
            {'field': None, 'label': "ФИО", 'type': QLineEdit, 'data': "name", 'access': ["Учитель", "Зав. кафедрой", "Методист", "dev"]},
            {'field': None, 'label': "Предмет", 'type': QLineEdit, 'data': "subject", 'access': ["Учитель", "Зав. кафедрой", "Методист", "dev"]},
            {'field': None, 'label': "Категория", 'type': QLineEdit, 'data': "qualification", 'access': ["Учитель", "Зав. кафедрой", "Методист", "dev"]},
            {'field': None, 'label': "Кафедра", 'type': QComboBox, 'data': "department", 'items': ['Кафедра точных наук', 'Кафедра естественных наук', 'Кафедра общественно-гуманитарных наук', 'Метод. объединение учителей физкультуры'], 'access': ["Учитель", "Зав. кафедрой", "Методист", "dev"]},
            {'field': None, 'label': "Доступ", 'type': QComboBox, 'data': "access", 'items': ['Учитель', 'Зав. кафедрой', 'Методист'], 'access': ["Методист", "dev"]}
        ]
        self.fields = []
        self.draw()
        if self.data:
            self.fill()

    def draw(self):
        layout = QVBoxLayout()
        for i in self.pattern:
            if self.root.access in i['access']:
                row = QHBoxLayout()
                lbl = QLabel(i['label'])
                row.addWidget(lbl, stretch=1)
                field = i['type']()
                if type(field) == QComboBox:
                    field.addItems(i['items'])
                i['field'] = field
                row.addWidget(field, stretch=3)
                layout.addLayout(row)
        layout.addWidget(QHLine())
        SBtn = QPushButton("Сохранить")
        SBtn.clicked.connect(self.updatePersona)
        ToggleBtn = QPushButton("Отключить профиль")
        ToggleBtn.clicked.connect(self.unTeach)
        PassBtn = QPushButton("Пароль")
        PassBtn.clicked.connect(self.seePass)
        btns = QHBoxLayout()
        btns.addWidget(SBtn, stretch=3)
        btns.addWidget(ToggleBtn, stretch=3)
        btns.addWidget(PassBtn, stretch=1)
        layout.addLayout(btns)
        self.setLayout(layout)
    
    def fill(self):
        for i in self.pattern:
            field = i['field']
            if type(field) == QLineEdit:
                print(i, ", ", self.data)
                field.setText(self.data[i['data']])
            elif type(field) == QComboBox:
                index = field.findText(self.data[i['data']])
                field.setCurrentIndex(index)

    def update(self, data: dict):
        self.data = data
        self.fill()

    def seePass(self):
        try: self.data['id']
        except KeyError:    #Если в словаре данных не находится индекс, прерываем функцию.
            warning = QMessageBox(self)
            warning.setWindowTitle("Ошибка")
            warning.setText("Профиль преподавателя не выбран!")
            warning.setIcon(QMessageBox.Warning)
            warning.exec_()
            return 0
        response = requests.get(server_host + "tPass", params={'id': self.data['id']})
        msg = QMessageBox(self)
        msg.setWindowTitle("Пароль")
        msg.setText(f"Пароль пользователя {self.data['name']}: \n{response.json()['password']}")
        msg.setIcon(QMessageBox.Information)
        msg.exec_()

    def updatePersona(self):
        try: self.data['id']
        except KeyError:    #Если в словаре данных не находится индекс, прерываем функцию.
            warning = QMessageBox(self)
            warning.setWindowTitle("Ошибка")
            warning.setText("Профиль преподавателя не выбран!")
            warning.setIcon(QMessageBox.Warning)
            warning.exec_()
            return 0
        vals = []
        for i in self.pattern:
            field = i['field']
            if type(field) == QLineEdit:
                vals.append(field.text())
            elif type(field) == QComboBox:
                vals.append(field.currentText())
        pack = {'id': self.data['id'], 'name': vals[0], 'subj': vals[1], 'qual': vals[2], 'dep': vals[3], 'access': vals[4]} #имя, предмет, категория, кафедра, доступ
        print(pack)
        request = requests.get(server_host + "reTeach", params=pack)
        print(request.text)
        self.root.name = vals[0]
        self.root.subj = vals[1]
        self.root.qual = vals[2]
        self.root.dep = vals[3]
        self.root.setWindowTitle(f"PlanTable - {self.root.name}: {self.root.subj} ({self.root.access})")

    def unTeach(self):
        warning = QMessageBox(self)
        warning.setWindowTitle("Отключение профиля")
        warning.setText("Отключенный профиль не будет отображаться в отчёте, но может быть включён в любое время. \nПродолжить?")
        warning.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
        warning.setIcon(QMessageBox.Question)
        button = warning.exec_()

        if button == QMessageBox.Yes:
            try: self.data['id']
            except KeyError:    #Если в словаре данных не находится индекс, прерываем функцию.
                warning = QMessageBox(self)
                warning.setWindowTitle("Ошибка")
                warning.setText("Профиль преподавателя не выбран!")
                warning.setIcon(QMessageBox.Warning)
                warning.exec_()
                return 0
            request = requests.get(server_host + "unTeach", params={"id": self.data['id']})
            print(request.json()['value'])
            if request.json()['value'] == True:
                self.data['button'].setStyleSheet('border: 1px solid black;\
                                background-color: #ffffff')
            else:
                self.data['button'].setStyleSheet('border: 1px solid black;\
                                background-color: #ababab')

        else:
            print("No!")
       

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("PlanTable")
        self.setMinimumSize(QSize(1024, 720))
        
        toolbar = QToolBar("My main toolbar")
        self.addToolBar(toolbar)

        self.teach = 0
        self.flag = "NVM"
        self.w = AuthDialog(self)
        self.w.exec_()
       
        if self.flag == "NVM":
            sys.exit(0)
        else:
            print(self.flag)
            self.name, self.teach, self.subj, self.access, self.qual, self.dep = self.flag.split(", ")
            print(self.teach, self.access)
            self.setWindowTitle(f"PlanTable - {self.name}: {self.subj} ({self.access})")
            self.pages = QStackedLayout()   
            self.tables = []
            self.draw()

    def draw(self): 
        self.page_main = QWidget()
        PL_main = QVBoxLayout()
        btn_plan = QPushButton("Таблицы")
        btn_plan.setMinimumHeight(75)
        btn_plan.clicked.connect(lambda: self.goTo(self.page_plan))
        PL_main.addWidget(btn_plan)
        btn_report = QPushButton("Документ")
        btn_report.clicked.connect(self.docx_gen)
        btn_report.setMinimumHeight(75)
        PL_main.addWidget(btn_report)
        btn_users = QPushButton("Данные преподавателя")
        btn_users.clicked.connect(lambda: self.goTo(self.page_users))
        btn_users.setMinimumHeight(75)
        PL_main.addWidget(btn_users)
        PL_main.setAlignment(Qt.AlignVCenter)
        self.page_main.setLayout(PL_main)
        self.pages.addWidget(self.page_main)

        self.page_users = QWidget()
        PL_users = QVBoxLayout()
        header = QHBoxLayout()
        backBtn = QPushButton("Назад")
        backBtn.clicked.connect(lambda: self.goTo(self.page_main))
        header.addWidget(backBtn, stretch=1)
        header.addWidget(self.lbl("Данные преподавателя"), stretch=6)
        PL_users.addLayout(header)
        self.PData = PDataForm(self)
        PL_users.addWidget(self.PData, stretch=1)

        OData = Color("#bababa")
        ODataLayout = QVBoxLayout()
        OData.setLayout(ODataLayout)
        table = QWidget()
        tRows = QVBoxLayout(table)
        response = requests.get(server_host + "tData", params={"id": self.teach})
        self.teachers = response.json()['teachers']
        for i in range(len(self.teachers)):
            teacher = self.teachers[i]
            btn = TableBtn()
            btn.index(i)
            if teacher['active'] == 1:
                color = "#ffffff"
            elif teacher['active'] == 0:
                color = "#ababab"
            btn.setStyleSheet(f'border: 1px solid black;\
                            background-color: {color}')
            btn.setMaximumHeight(50)
            btn.setMinimumHeight(50)
            btnLayout = QHBoxLayout(btn)
            lbl = QLabel(teacher['name'])
            lbl.setStyleSheet('border: 1px solid black;')
            btnLayout.addWidget(lbl, stretch=3)
            lbl = QLabel(teacher['department'])
            lbl.setStyleSheet('border: 1px solid black;')
            btnLayout.addWidget(lbl, stretch=5)
            lbl = QLabel(teacher['subject'])
            lbl.setStyleSheet('border: 1px solid black;')
            btnLayout.addWidget(lbl, stretch=3)
            btn.setSizePolicy(QSizePolicy.Preferred, QSizePolicy.Preferred)
            btn.clicked.connect(self.PDataUpdate)
            tRows.addWidget(btn)
        scroller = QScrollArea()
        scroller.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOn)
        scroller.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroller.setWidgetResizable(True)                           
        scroller.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        scroller.setAlignment(Qt.AlignTop | Qt.AlignLeft)
        scroller.setWidget(table)
        ODataLayout.addWidget(scroller)
        PL_users.addWidget(OData, stretch=1)

        self.page_users.setLayout(PL_users)
        self.pages.addWidget(self.page_users)

        self.page_plan = QWidget()
        PL_plan = QVBoxLayout()
        PL_plan.addWidget(self.lbl("Разделы документа"))
        btn_plan_01 = QPushButton("Учебно-методическая работа")
        btn_plan_01.setMinimumHeight(75)
        btn_plan_01.clicked.connect(lambda: self.goTo(self.page_plan_p1))
        PL_plan.addWidget(btn_plan_01)
        btn_plan_02 = QPushButton("Научно-методическая работа")
        btn_plan_02.setMinimumHeight(75)
        btn_plan_02.clicked.connect(lambda: self.goTo(self.page_plan_p2))
        PL_plan.addWidget(btn_plan_02)
        btn_plan_03 = QPushButton("Повышение квалификации")
        btn_plan_03.setMinimumHeight(75)
        btn_plan_03.clicked.connect(lambda: self.goTo(self.page_plan_p3))
        PL_plan.addWidget(btn_plan_03)
        btn_plan_04 = QPushButton("Работа с обучающимися")
        btn_plan_04.setMinimumHeight(75)
        btn_plan_04.clicked.connect(lambda: self.goTo(self.page_plan_p4))
        PL_plan.addWidget(btn_plan_04)
        PL_plan.addLayout(self.sml_nav())
        PL_plan.setAlignment(Qt.AlignVCenter)
        self.page_plan.setLayout(PL_plan)
        self.pages.addWidget(self.page_plan)

        self.page_plan_p1 = QWidget()
        PL_plan_p1 = QVBoxLayout()
        PL_plan_p1.addWidget(self.lbl("Учебно-методическая Работа"))
        btn_table1 = QPushButton("Проведение открытых уроков, классных часов, предметных недель, других мероприятий")
        btn_table1.setMinimumHeight(75)
        btn_table1.clicked.connect(lambda: self.goTo(self.tables[0].widget))
        PL_plan_p1.addWidget(btn_table1)
        btn_table2 = QPushButton("Участие в подготовке и проведении лицейских мероприятий")
        btn_table2.setMinimumHeight(75)
        btn_table2.clicked.connect(lambda: self.goTo(self.tables[1].widget))
        PL_plan_p1.addWidget(btn_table2)
        PL_plan_p1.addLayout(self.nav())
        self.page_plan_p1.setLayout(PL_plan_p1)
        self.pages.addWidget(self.page_plan_p1)

        self.tables.append(Table(self, "t-01"))
        self.pages.addWidget(self.tables[-1].widget)
        self.tables[0].btn_prev.setEnabled(False)
        
        self.tables.append(Table(self, "t-02"))
        self.pages.addWidget(self.tables[-1].widget)

        self.page_plan_p2 = QWidget()
        PL_plan_p2 = QVBoxLayout()
        PL_plan_p2.addWidget(self.lbl("Научно-методическая Работа"))
        btn_table1 = QPushButton("Запланированные Мероприятия")
        btn_table1.setMinimumHeight(75)
        btn_table1.clicked.connect(lambda: self.goTo(self.tables[2].widget))
        PL_plan_p2.addWidget(btn_table1)
        btn_table2 = QPushButton("Работа в рамках творческих групп, инновационной/стажировочной деятельности площадок")
        btn_table2.setMinimumHeight(75)
        btn_table2.clicked.connect(lambda: self.goTo(self.tables[3].widget))
        PL_plan_p2.addWidget(btn_table2)
        btn_table3 = QPushButton("Экспертная Деятельность")
        btn_table3.setMinimumHeight(75)
        btn_table3.clicked.connect(lambda: self.goTo(self.tables[4].widget))
        PL_plan_p2.addWidget(btn_table3)
        PL_plan_p2.addLayout(self.nav())
        self.page_plan_p2.setLayout(PL_plan_p2)
        self.pages.addWidget(self.page_plan_p2)

        self.tables.append(Table(self, "t-03"))
        self.pages.addWidget(self.tables[-1].widget)

        self.tables.append(Table(self, "t-04"))
        self.pages.addWidget(self.tables[-1].widget)

        self.tables.append(Table(self, "t-05"))
        self.pages.addWidget(self.tables[-1].widget)

        self.page_plan_p3 = QWidget()
        PL_plan_p3 = QVBoxLayout()
        PL_plan_p3.addWidget(self.lbl("Повышение Квалификации"))
        btn_table1 = QPushButton("Обучение на курсах повышения квалификации, посещение опорных школ и др.")
        btn_table1.setMinimumHeight(75)
        btn_table1.clicked.connect(lambda: self.goTo(self.tables[5].widget))
        PL_plan_p3.addWidget(btn_table1)
        btn_table2 = QPushButton("Участие в сертифицированные вебинарах, семинарах и др.")
        btn_table2.setMinimumHeight(75)
        btn_table2.clicked.connect(lambda: self.goTo(self.tables[6].widget))
        PL_plan_p3.addWidget(btn_table2)
        btn_table3 = QPushButton("Участие в конкурсах профессионального мастерства")
        btn_table3.setMinimumHeight(75)
        btn_table3.clicked.connect(lambda: self.goTo(self.tables[7].widget))
        PL_plan_p3.addWidget(btn_table3)
        btn_table4 = QPushButton("Обобщение и представление опыта работы")
        btn_table4.setMinimumHeight(75)
        btn_table4.clicked.connect(lambda: self.goTo(self.tables[8].widget))
        PL_plan_p3.addWidget(btn_table4)
        btn_table5 = QPushButton("Участие в диагностике профессиональных дефицитов/предметных компетенций")
        btn_table5.setMinimumHeight(75)
        btn_table5.clicked.connect(lambda: self.goTo(self.tables[9].widget))
        PL_plan_p3.addWidget(btn_table5)
        btn_table6 = QPushButton("Участие во внешкольных мероприятий")
        btn_table6.setMinimumHeight(75)
        btn_table6.clicked.connect(lambda: self.goTo(self.tables[10].widget))
        PL_plan_p3.addWidget(btn_table6)
        btn_table7 = QPushButton("Посещение уроков, кл.часов, мероприятий у коллег в школе")
        btn_table7.setMinimumHeight(75)
        btn_table7.clicked.connect(lambda: self.goTo(self.tables[11].widget))
        PL_plan_p3.addWidget(btn_table7)
        PL_plan_p3.addLayout(self.nav())
        self.page_plan_p3.setLayout(PL_plan_p3)
        self.pages.addWidget(self.page_plan_p3)

        self.tables.append(Table(self, "t-06"))
        self.pages.addWidget(self.tables[-1].widget)

        self.tables.append(Table(self, "t-07"))
        self.pages.addWidget(self.tables[-1].widget)

        self.tables.append(Table(self, "t-08"))
        self.pages.addWidget(self.tables[-1].widget)

        self.tables.append(Table(self, "t-09"))
        self.pages.addWidget(self.tables[-1].widget)

        self.tables.append(Table(self, "t-10"))
        self.pages.addWidget(self.tables[-1].widget)

        self.tables.append(Table(self, "t-11"))
        self.pages.addWidget(self.tables[-1].widget)

        self.tables.append(Table(self, "t-12"))
        self.pages.addWidget(self.tables[-1].widget)

        self.page_plan_p4 = QWidget()
        PL_plan_p4 = QVBoxLayout()
        PL_plan_p4.addWidget(self.lbl("Работа с обучающимися"))
        btn_table1 = QPushButton("Участие обучающихся в конкурсных мероприятиях, входящих в перечень, \n утвержденный приказом Министертсва науки и высшего образования РФ")
        btn_table1.setMinimumHeight(75)
        btn_table1.clicked.connect(lambda: self.goTo(self.tables[12].widget))
        PL_plan_p4.addWidget(btn_table1)
        btn_table2 = QPushButton("Участие обучающихся в других конкурсных мероприятиях, научно-практических конференциях, ШРД, ФНР и др.")
        btn_table2.setMinimumHeight(75)
        btn_table2.clicked.connect(lambda: self.goTo(self.tables[13].widget))
        PL_plan_p4.addWidget(btn_table2)
        btn_table3 = QPushButton("Участие обучающихся в соревнованиях профессиональных компетенций ")
        btn_table3.setMinimumHeight(75)
        btn_table3.clicked.connect(lambda: self.goTo(self.tables[14].widget))
        PL_plan_p4.addWidget(btn_table3)
        btn_table4 = QPushButton("Дополнительные общеразвивающие программы (ДОП) по подготовке обучющихся 9-11 классов к ВсОШ")
        btn_table4.setMinimumHeight(75)
        btn_table4.clicked.connect(lambda: self.goTo(self.tables[15].widget))
        PL_plan_p4.addWidget(btn_table4)
        btn_table5 = QPushButton("Участие в профильных сменах")
        btn_table5.setMinimumHeight(75)
        btn_table5.clicked.connect(lambda: self.goTo(self.tables[16].widget))
        PL_plan_p4.addWidget(btn_table5) 
        PL_plan_p4.addLayout(self.nav())
        self.page_plan_p4.setLayout(PL_plan_p4)
        self.pages.addWidget(self.page_plan_p4)

        self.tables.append(Table(self, "t-13"))
        self.pages.addWidget(self.tables[-1].widget)

        self.tables.append(Table(self, "t-14"))
        self.pages.addWidget(self.tables[-1].widget)

        self.tables.append(Table(self, "t-15"))
        self.pages.addWidget(self.tables[-1].widget)

        self.tables.append(Table(self, "t-16"))
        self.pages.addWidget(self.tables[-1].widget)

        self.tables.append(Table(self, "t-17"))
        self.tables[-1].Btn_next.setEnabled(False)
        self.pages.addWidget(self.tables[-1].widget)

        widget = QWidget()
        widget.setLayout(self.pages)
        self.setCentralWidget(widget)

    def PDataUpdate(self):
        sender = self.sender()
        data = self.teachers[sender.id]
        data['button'] = sender
        self.PData.update(data)

    def goTo(self, page):
        self.pages.setCurrentWidget(page)
    
    def next_plan_table(self, curtable):
        widget = self.tables[self.tables.index(curtable)+1].widget
        self.pages.setCurrentWidget(widget)

    def prev_plan_table(self, curtable):
        widget = self.tables[self.tables.index(curtable)-1].widget
        self.pages.setCurrentWidget(widget)

    def button_pushed(self):
        print("Кнопка работает, начальник! Честно-честно!")

    def sml_nav(self):
        nav = QHBoxLayout()
        btn_nav_menu = QPushButton("В меню")
        btn_nav_menu.setMinimumHeight(75)
        btn_nav_menu.clicked.connect(lambda: self.goTo(self.page_main))
        nav.addWidget(btn_nav_menu)
        btn_nav_doc = QPushButton("К созданию документа")
        btn_nav_doc.setMinimumHeight(75)
        btn_nav_doc.clicked.connect(self.docx_gen)
        nav.addWidget(btn_nav_doc)
        return nav

    def nav(self):
        nav = QHBoxLayout()
        btn_01 = QPushButton("В Меню")
        btn_01.clicked.connect(lambda: self.goTo(self.page_main))
        nav.addWidget(btn_01)
        nav.setAlignment(Qt.AlignBottom)
        return nav

    def lbl(self, text):
        lbl = QLabel(text)
        font = lbl.font()
        font.setPointSize(18)
        lbl.setFont(font)
        lbl.setAlignment(Qt.AlignHCenter | Qt.AlignTop)
        return lbl
    
    def initials(self, name):
        fname = name.split(" ")
        if len(fname) > 2:
            tSurname, tName, tPatronimic = fname
        else:
            tSurname = fname[0]
            tName = "John"
            tPatronimic = "Doe"
        return str(tSurname) + " " + str(tName)[0] + "." + str(tPatronimic)[0] + "."

    def docx_gen(self):
        dial = DocDialog(self)
        res = dial.exec_()

        if res == 0:
            return 0
        elif res == 2 or res == 3:
            response = requests.get(server_host + "docData", params={"user": self.teach, "docType": res})
            print(self.dep, response.json()["teachers"])
            trs = response.json()['teachers']
            selfEd_themes = response.json()['selfEd']
            dep = self.dep
            if "Кафедра" in dep:
                dep = "Кафедры" + dep[7:]
            elif "Метод." in dep:
                dep = "Метод. объединения" + dep[18:]
    
        print("Генерация документа начата...")
        doc = Document()
        self.style = doc.styles["Normal"]
        self.style.font.name = "Times New Roman"
        self.style.font.size = Pt(12)
        doc.sections[0].orientation = WD_ORIENT.LANDSCAPE
        doc.sections[0].page_width = Mm(297)
        doc.sections[0].page_height = Cm(21)
        doc.sections[0].left_margin = Mm(30)
        doc.sections[0].right_margin = Mm(15)
        doc.sections[0].top_margin = Mm(20)
        doc.sections[0].bottom_margin = Mm(20)
        curdate = datetime.now()
        if curdate.month > 9:
            years = (curdate.year, curdate.year+1)
        else:
            years = (curdate.year-1, curdate.year)
        if res == 1: 
            head = doc.add_paragraph()
            head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hline = head.add_run('Индивидуальный план работы учителя МБОУ "Лицей №1" г.Братска')
            hline.bold = True
            hline.font.size = Pt(14)
        elif res == 2:
            lines = ["Анализ работы", dep]
            for line in lines:
                head = doc.add_paragraph()
                head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                hline = head.add_run(line)
                hline.bold = True
                hline.font.size = Pt(14)
        elif res == 3:
            lines = ["Анализ работы", "Лицейских преподавателей (временный заголовок)"]
            for line in lines:
                head = doc.add_paragraph()
                head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                hline = head.add_run(line)
                hline.bold = True
                hline.font.size = Pt(14)
        

        year = doc.add_paragraph()
        year.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        a = year.add_run(str(years[0]))
        a.bold = True
        a.font.size = Pt(14)
        a = year.add_run(" — ")
        a.font.size = Pt(14)
        a = year.add_run(str(years[1]))
        a.bold = True
        a.font.size = Pt(14)
        year.add_run(" учебный год")

        if res == 1:
            teachData = doc.add_paragraph()
            teachData.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            t_text = self.initials(self.name)
            if len(t_text) < 21:
                diff = 21 - len(t_text)
                pref = "_" * (diff//2)
                posf = "_" * (diff - (diff//2))
                t_text = pref + t_text + posf
                print(t_text)
            teach = teachData.add_run(t_text)
            teach.font.size = Pt(12)
            teach = teachData.add_run("     ")
            teach.font.size = Pt(12)
            t_text = self.subj
            if len(t_text) < 27:
                diff = 27 - len(t_text)
                pref = "_" * (diff//2)
                posf = "_" * (diff - (diff//2))
                t_text = pref + t_text + posf
                print(t_text)
            teach = teachData.add_run(t_text)
            teach.font.size = Pt(12)
            teach = teachData.add_run(", квалификационная категория")
            teach.font.size = Pt(12)
            t_text = self.qual
            if len(t_text) < 21:
                diff = 21 - len(t_text)
                pref = "_" * (diff//2)
                posf = "_" * (diff - (diff//2))
                t_text = pref + t_text + posf
                print(t_text)
            teach = teachData.add_run(t_text)
            teach.font.size = Pt(12)

        index = -1
        for block in docTables:
            doc.add_paragraph()
            blockhead = doc.add_paragraph(style="List Number")
            a = blockhead.add_run(block[0])
            a.bold = True
            a.italic = True
            a.font.size = Pt(14)
            for table in block[1:]:
                index += 1
                data = self.tables[index].data
                print(self.tables[index].name, end=" ")
                
                if index != 2:
                    p = doc.add_paragraph()
                    p.add_run(table[0]).bold = True
                else:
                    if res == 1:
                        theme = doc.add_paragraph()
                        htheme = theme.add_run("Тема самообразования: ")
                        htheme.bold = True
                        theme.add_run(self.tables[2].theme.text())

                        stage = doc.add_paragraph()
                        hstage = stage.add_run("Сроки работы над темой: ")
                        hstage.bold = True
                        stage.add_run(self.tables[2].stage.text())

                        method = doc.add_paragraph()
                        hmethod = method.add_run("Образовательная технология/Метод обучения: ")
                        hmethod.bold = True
                        method.add_run(self.tables[2].method.text())
                    elif res == 2 or res == 3:
                        themes = doc.add_table(rows = 1, cols = 5)
                        themes.style = "Table Grid"
                        headline = themes.rows[0].cells
                        headline[0].text = "№"
                        headlines = ["Преподаватель", "Тема", "Сроки", "Метод обучения"]
                        for i in range(1, len(headline)):
                            headline[i].text = headlines[i-1]
                        make_rows_bold(themes.rows[0])
                        for teacher_num in range(len(trs)):
                            dataLine = [teacher_num + 1, trs[teacher_num][1]] + selfEd_themes[teacher_num]
                            print(dataLine)
                            row = themes.add_row().cells
                            for cell_num in range(len(row)):
                                print(type(row), type(dataLine))
                                row[cell_num].text = str(dataLine[cell_num])

                        col = themes.columns[0]
                        col.width=Cm(1.5)

                        doc.add_paragraph()


                tab = doc.add_table(rows = 1, cols = len(table[1])+1)
                tab.style = "Table Grid"
                hedline = tab.rows[0].cells
                hedline[0].text = "№"
                for i in range(1, len(hedline)):
                    hedline[i].text = table[1][i-1]

                col = tab.columns[0]
                col.width=Cm(1.5)

                make_rows_bold(tab.rows[0])
                cnt = 1
                if res == 1:
                    print("INDIVIDUAL")
                    for line in data:
                        if int(line["teacher"]) != int(self.teach):
                            continue
                        print(line)
                        row = tab.add_row().cells
                        row[0].text = str(cnt)
                        cnt += 1
                        for i in range(0, len(table[2])):
                            num = table[2][i]
                            if index == 15 and (num == "date1" or num == 9):
                                row[i+1].text = f"{line[num]} - {line[num+1]}"
                            else:
                                if line[num] == True or line[num] == "True":
                                    row[i+1].text = "+"
                                elif line[num] == False or line[num] == "False":
                                    row[i+1].text = "-"
                                else:
                                    row[i+1].text = line[num]
                elif res == 2 or res == 3:
                    if res == 2:
                        print("DEPARTMENT")
                    elif res == 3:
                        print("GENERAL")
                    for teacher in trs:
                        print(teacher)
                        row = tab.add_row().cells
                        row[0].merge(row[-1])
                        row[0].text = teacher[1]
                        make_row_bold(row)
                        for line in data:
                            if int(line["teacher"]) != int(teacher[0]):
                                continue
                            print(line)
                            row = tab.add_row().cells
                            row[0].text = str(cnt)
                            cnt += 1
                            for i in range(0, len(table[2])):
                                num = table[2][i]
                                if index == 15 and (num == "date1" or num == 9):
                                    row[i+1].text = f"{line[num]} - {line[num+1]}"
                                else:
                                    if line[num] == True or line[num] == "True":
                                        row[i+1].text = "+"
                                    elif line[num] == False or line[num] == "False":
                                        row[i+1].text = "-"
                                    elif num == "student":
                                        row[i+1].text = str(line[num]) + ", " + str(line["studClass"])
                                    else:
                                        row[i+1].text = line[num]
                if index == 15:
                    break
                doc.add_paragraph()
        

       
        if res == 3:
            doc.save(f'Lyceum_GenDoc.docx')
            print(f'Документ "Lyceum_GenDoc.docx" готов!')
        elif res == 2:
            doc.save(f'{self.dep}_GenDoc.docx')
            print(f'Документ "{self.dep}_GenDoc.docx" готов!')
        elif res == 1:
            doc.save(f"{self.name}_IndDoc.docx")
            print(f'Документ "{self.name}_IndDoc.docx" готов!')
        
        

if __name__ == "__main__":
    app = QApplication(sys.argv)

    font = app.font()
    font.setPointSize(12)
    app.setFont(font)

    window = MainWindow()
    window.show()

    app.exec_()